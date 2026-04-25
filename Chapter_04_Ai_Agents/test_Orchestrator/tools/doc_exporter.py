"""
tools/doc_exporter.py
Layer 3 — Export test plan markdown to Word (.docx) format.
No LLM calls.
"""

import re
from pathlib import Path

TMP_DIR = Path(__file__).parent.parent / ".tmp"


def export_to_doc(content: str, project_key: str, timestamp: str) -> Path:
    """Convert markdown string to a styled .docx Word document."""
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    doc_path = TMP_DIR / f"test_plan_{project_key}_{timestamp}.docx"
    _build_docx(content, doc_path)
    return doc_path


def _build_docx(content: str, output_path: Path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Style helpers
    def set_heading(para, level: int):
        colors = {1: RGBColor(0x1a, 0x1a, 0x2e), 2: RGBColor(0x16, 0x21, 0x3e), 3: RGBColor(0x0f, 0x34, 0x60)}
        sizes = {1: 20, 2: 16, 3: 13}
        run = para.runs[0] if para.runs else para.add_run(para.text)
        run.bold = True
        run.font.size = Pt(sizes.get(level, 12))
        run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))

    lines = content.split("\n")
    i = 0
    table_rows = []
    last_was_list = False  # track whether the previous content was a list item

    while i < len(lines):
        line = lines[i]

        # Heading 1
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(line[2:].strip(), level=1)
            set_heading(p, 1)
            last_was_list = False

        # Heading 2
        elif line.startswith("## ") and not line.startswith("### "):
            p = doc.add_heading(line[3:].strip(), level=2)
            set_heading(p, 2)
            last_was_list = False

        # Heading 3
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            set_heading(p, 3)
            last_was_list = False

        # Table row — collect and render as a table
        elif line.startswith("|"):
            table_rows.append(line)
            while i + 1 < len(lines) and lines[i + 1].startswith("|"):
                i += 1
                table_rows.append(lines[i])
            _add_table(doc, table_rows)
            table_rows = []
            last_was_list = False

        # Bullet list
        elif line.startswith("- ") or line.startswith("* "):
            text = _strip_md(line[2:].strip())
            p = doc.add_paragraph(text, style="List Bullet")
            p.runs[0].font.size = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            last_was_list = True

        # Numbered list
        elif re.match(r"^\d+\.\s", line):
            text = _strip_md(re.sub(r"^\d+\.\s", "", line))
            p = doc.add_paragraph(text, style="List Number")
            p.runs[0].font.size = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            last_was_list = True

        # Blockquote
        elif line.startswith("> "):
            p = doc.add_paragraph(_strip_md(line[2:]))
            p.paragraph_format.left_indent = Inches(0.4)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                run.font.size = Pt(10)
            last_was_list = False

        # Empty line — skip if between/after list items to avoid huge gaps
        elif line.strip() == "":
            next_line = lines[i + 1] if i + 1 < len(lines) else ""
            next_is_list = next_line.startswith("- ") or next_line.startswith("* ")
            if not last_was_list and not next_is_list:
                doc.add_paragraph("")

        # Normal paragraph
        else:
            text = _strip_md(line)
            if text.strip():
                p = doc.add_paragraph()
                _add_formatted_run(p, line)
            last_was_list = False

        i += 1

    doc.save(str(output_path))


def _add_table(doc, rows: list[str]):
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Filter separator rows (---|---|---)
    data_rows = [r for r in rows if not re.match(r"^\|[\s\-:]+\|", r)]
    if not data_rows:
        return

    parsed = []
    for row in data_rows:
        # Strip leading/trailing pipes then split; handle edge whitespace
        stripped = row.strip()
        if stripped.startswith("|"):
            stripped = stripped[1:]
        if stripped.endswith("|"):
            stripped = stripped[:-1]
        cells = [c.strip() for c in stripped.split("|")]
        parsed.append(cells)

    if not parsed:
        return

    num_cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=num_cols)
    table.style = "Table Grid"

    # Distribute column widths: give Summary column more space
    available_width = 6.0  # inches (page width minus margins)
    if num_cols == 5:
        col_widths = [1.0, 0.8, 2.4, 0.9, 0.9]
    else:
        col_widths = [available_width / num_cols] * num_cols

    for c_idx, width in enumerate(col_widths[:num_cols]):
        for cell in table.columns[c_idx].cells:
            cell.width = Inches(width)

    for r_idx, row in enumerate(parsed):
        for c_idx in range(num_cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = _strip_md(cell_text)
            para = cell.paragraphs[0]
            run = para.runs[0] if para.runs else para.add_run(cell.text)
            run.font.size = Pt(9)
            if r_idx == 0:
                run.bold = True
                # Header row background: light blue-grey
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D9E1F2")
                tcPr.append(shd)


def _strip_md(text: str) -> str:
    """Strip markdown formatting for plain Word text."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


def _add_formatted_run(para, line: str):
    """Add a paragraph run preserving bold/italic from markdown."""
    from docx.shared import Pt
    # Split on **bold** and *italic* markers
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", line)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            para.add_run(part)
    for run in para.runs:
        run.font.size = Pt(10)
