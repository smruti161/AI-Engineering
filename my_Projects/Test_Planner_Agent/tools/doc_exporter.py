"""
tools/doc_exporter.py
Layer 3 — Export test plan markdown to Word (.docx) format.
No LLM calls.
"""

import re
from pathlib import Path

TMP_DIR = Path(__file__).parent.parent / ".tmp"


def export_to_doc(content: str, project_key: str, timestamp: str) -> Path:
    """Convert markdown string to a styled .docx Word document."""
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    doc_path = TMP_DIR / f"test_plan_{project_key}_{timestamp}.docx"
    _build_docx(content, doc_path)
    return doc_path


def _build_docx(content: str, output_path: Path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Style helpers
    def set_heading(para, level: int):
        colors = {1: RGBColor(0x1a, 0x1a, 0x2e), 2: RGBColor(0x16, 0x21, 0x3e), 3: RGBColor(0x0f, 0x34, 0x60)}
        sizes = {1: 20, 2: 16, 3: 13}
        run = para.runs[0] if para.runs else para.add_run(para.text)
        run.bold = True
        run.font.size = Pt(sizes.get(level, 12))
        run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))

    lines = content.split("\n")
    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Heading 1
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(line[2:].strip(), level=1)
            set_heading(p, 1)

        # Heading 2
        elif line.startswith("## ") and not line.startswith("### "):
            p = doc.add_heading(line[3:].strip(), level=2)
            set_heading(p, 2)

        # Heading 3
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            set_heading(p, 3)

        # Table row — collect and render as a table
        elif line.startswith("|"):
            table_rows.append(line)
            # Look ahead to collect all table rows
            while i + 1 < len(lines) and lines[i + 1].startswith("|"):
                i += 1
                table_rows.append(lines[i])
            _add_table(doc, table_rows)
            table_rows = []

        # Bullet list
        elif line.startswith("- ") or line.startswith("* "):
            text = _strip_md(line[2:].strip())
            p = doc.add_paragraph(text, style="List Bullet")
            p.runs[0].font.size = Pt(10)

        # Numbered list
        elif re.match(r"^\d+\.\s", line):
            text = _strip_md(re.sub(r"^\d+\.\s", "", line))
            p = doc.add_paragraph(text, style="List Number")
            p.runs[0].font.size = Pt(10)

        # Blockquote
        elif line.startswith("> "):
            p = doc.add_paragraph(_strip_md(line[2:]))
            p.paragraph_format.left_indent = Inches(0.4)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                run.font.size = Pt(10)

        # Empty line
        elif line.strip() == "":
            doc.add_paragraph("")

        # Normal paragraph
        else:
            text = _strip_md(line)
            if text.strip():
                p = doc.add_paragraph()
                _add_formatted_run(p, line)

        i += 1

    doc.save(str(output_path))


def _add_table(doc, rows: list[str]):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Filter separator rows (---|---|---)
    data_rows = [r for r in rows if not re.match(r"^\|[\s\-:]+\|", r)]
    if not data_rows:
        return

    parsed = []
    for row in data_rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    if not parsed:
        return

    num_cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=num_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(parsed):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= num_cols:
                break
            cell = table.cell(r_idx, c_idx)
            cell.text = _strip_md(cell_text)
            run = cell.paragraphs[0].runs
            if run:
                run[0].font.size = Pt(9)
                if r_idx == 0:
                    run[0].bold = True


def _strip_md(text: str) -> str:
    """Strip markdown formatting for plain Word text."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


def _add_formatted_run(para, line: str):
    """Add a paragraph run preserving bold/italic from markdown."""
    from docx.shared import Pt
    # Split on **bold** and *italic* markers
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", line)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            para.add_run(part)
    for run in para.runs:
        run.font.size = Pt(10)
